import concurrent.futures
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document  # pip3 install python-docx --break-system-packages


def extraer_texto_docx(ruta):
    try:
        doc = Document(str(ruta))
        partes = [p.text for p in doc.paragraphs]

        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    partes.append(celda.text)

        return "\n".join(partes)
    except Exception:
        return None


def _procesar_lote(lote, timeout):
    """
    Convierte un lote de archivos .doc copiandolos a una carpeta temporal
    propia y corriendo textutil ahi adentro (sin -output: por defecto
    genera el .txt al lado de cada archivo, en la misma carpeta temporal,
    nunca en el Drive real). Devuelve {str(ruta_original): texto_o_None}.
    """
    resultados_lote = {}

    with tempfile.TemporaryDirectory() as tmp:

        tmp_path = Path(tmp)
        copias = {}  # nombre temporal -> ruta original

        for idx, ruta in enumerate(lote):
            nombre_temp = f"{idx:03d}.doc"
            try:
                shutil.copy2(ruta, tmp_path / nombre_temp)
                copias[nombre_temp] = ruta
            except (OSError, shutil.Error):
                resultados_lote[str(ruta)] = None

        if copias:

            args = ["textutil", "-convert", "txt"] + [str(tmp_path / n) for n in copias]

            try:
                subprocess.run(args, capture_output=True, text=True, timeout=timeout)
            except subprocess.TimeoutExpired:
                pass

            for nombre_temp, ruta_original in copias.items():
                salida = tmp_path / (Path(nombre_temp).stem + ".txt")
                try:
                    resultados_lote[str(ruta_original)] = salida.read_text(encoding="utf-8", errors="ignore")
                except (FileNotFoundError, OSError):
                    resultados_lote[str(ruta_original)] = None

    return resultados_lote


def extraer_textos_doc_en_lote(rutas, tamano_lote=40, timeout=90, procesos_simultaneos=6, cancelar_evento=None):
    """
    Convierte muchos .doc (formato binario viejo de Word) usando textutil,
    en lotes de "tamano_lote", corriendo varios lotes AL MISMO TIEMPO
    (procesos_simultaneos) para aprovechar los varios nucleos de la Mac
    en vez de esperar un lote atras del otro.

    Si se pasa cancelar_evento (threading.Event) y se activa a mitad de
    camino, se cancelan los lotes que todavia no arrancaron (los que ya
    estaban corriendo terminan igual, son come mucho unos pocos mas).

    Devuelve un dict {str(ruta_original): texto}. texto es None si ese
    archivo en particular no se pudo copiar o convertir.
    """
    resultados = {}

    lotes = [rutas[i:i + tamano_lote] for i in range(0, len(rutas), tamano_lote)]
    total_lotes = len(lotes)

    print(f"      {total_lotes} lotes, {procesos_simultaneos} en simultaneo...")

    completados = 0

    with concurrent.futures.ThreadPoolExecutor(max_workers=procesos_simultaneos) as executor:

        futuros = {executor.submit(_procesar_lote, lote, timeout): lote for lote in lotes}

        for futuro in concurrent.futures.as_completed(futuros):

            if cancelar_evento is not None and cancelar_evento.is_set():
                print("      Cancelado: no se esperan mas lotes (los que ya estaban corriendo, terminan).")
                for f in futuros:
                    f.cancel()
                break

            resultados.update(futuro.result())
            completados += 1

            if completados % 5 == 0 or completados == total_lotes:
                print(f"      {completados}/{total_lotes} lotes listos")

    return resultados
