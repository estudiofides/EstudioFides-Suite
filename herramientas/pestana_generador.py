"""
Pestaña "Generador Legal" para la app unificada (hub_app.py).
Misma lógica de herramientas/generador_legal.py. El original era una
clase que heredaba de tk.Tk (una ventana propia); acá hereda de
tk.Frame para poder vivir embebida dentro de una pestaña.
"""
import threading
import tkinter as tk
from pathlib import Path
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from src.tooltip import agregar_tooltip
from src.clientes import obtener_clientes
from src.ficha_cliente import leer_ficha

# ------------------------------------------------------------------
# Membrete (encabezado) y tamaño de hoja. Esto es a propósito lo único
# "de diseño" que hay separado arriba del archivo: son placeholders --
# Leandro va a querer cambiar el texto exacto (y a futuro, agregar el
# logo del estudio como imagen en vez de solo texto). Cambiarlo acá
# alcanza, no hace falta tocar generar_word().
# ------------------------------------------------------------------

MEMBRETE_NOMBRE = "ESTUDIO FIDES"
MEMBRETE_DETALLE = "Abogados y Escribanos · Rosario, Santa Fe"

# Carta Documento y Telegrama van en hoja tamaño Legal/Oficio (más
# larga que A4), como pidió Leandro. Los escritos y el contrato quedan
# en el tamaño por defecto de Word (A4), que es lo normal para
# presentar en un expediente.
_HOJA_LEGAL = {"Carta Documento (Intimación)", "Telegrama Laboral"}


def _agregar_membrete(doc):
    """Encabezado simple (nombre del estudio + detalle) arriba de cada
    documento generado. Placeholder de texto por ahora -- cuando
    Leandro tenga el membrete definitivo (y el logo, si quiere uno),
    se reemplaza esto por una imagen con doc.sections[0].header
    .paragraphs[0].add_run().add_picture(...)."""

    header = doc.sections[0].header
    p_nombre = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nombre = p_nombre.add_run(MEMBRETE_NOMBRE)
    run_nombre.bold = True
    run_nombre.font.size = Pt(13)

    p_detalle = header.add_paragraph()
    p_detalle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_detalle = p_detalle.add_run(MEMBRETE_DETALLE)
    run_detalle.font.size = Pt(9)


# Para el botón "Completar con datos de un cliente": por cada
# plantilla, qué campo(s) representan a "una persona" -- ahí puede ir
# el cliente elegido en el buscador. Cada rol dice en qué campo va el
# nombre, y opcionalmente en qué otro campo van el DNI/CUIT y el
# domicilio (si la plantilla los tiene por separado). Cuando no hay un
# campo de DNI separado pero el campo de nombre dice "y DNI" en su
# etiqueta, se pega el DNI al final del nombre (combinar_dni_en_nombre).
ROLES_PLANTILLA = {
    "Carta Documento (Intimación)": [
        {"campo_nombre": "Remitente (Nombre y DNI)", "combinar_dni_en_nombre": True},
        {"campo_nombre": "Destinatario", "campo_domicilio": "Domicilio Destinatario"},
    ],
    "Telegrama Laboral": [
        {"campo_nombre": "Remitente", "campo_dni": "CUIL Remitente"},
        {"campo_nombre": "Destinatario (Empleador)", "campo_domicilio": "Domicilio Laboral"},
    ],
    "Escrito: Conformidad (Sucesión)": [
        {"campo_nombre": "Nombre del Cliente", "campo_dni": "DNI/CUIT del Cliente"},
    ],
    "Contrato de Locación": [
        {"campo_nombre": "Propietario", "campo_dni": "DNI Propietario"},
        {"campo_nombre": "Inquilino", "campo_dni": "DNI Inquilino"},
    ],
}


PLANTILLAS = {
    "Carta Documento (Intimación)": [
        "Remitente (Nombre y DNI)", "Destinatario", "Domicilio Destinatario", "Localidad", "Texto de la intimación"
    ],
    "Telegrama Laboral": [
        "Remitente", "CUIL Remitente", "Destinatario (Empleador)", "Domicilio Laboral", "Texto del telegrama"
    ],
    "Escrito: Conformidad (Sucesión)": [
        "Nombre del Cliente", "DNI/CUIT del Cliente", "Carátula del Expediente", "Número de Expedte.", "Radicación (Ej: Juzgado Civil y Com. 1 - Rosario)"
    ],
    "Contrato de Locación": [
        "Propietario", "DNI Propietario", "Inquilino", "DNI Inquilino", "Dirección del Inmueble",
        "Precio Mensual ($)", "Plazo en meses", "Garantías (Detalle)", "Servicios a cargo de", "Inventario y estado"
    ]
}


def generar_word(tipo_doc, datos, ruta_guardado):
    doc = Document()
    estilo = doc.styles['Normal']
    estilo.font.name = 'Arial'
    estilo.font.size = Pt(11)

    if tipo_doc in _HOJA_LEGAL:
        seccion = doc.sections[0]
        seccion.page_width = Inches(8.5)
        seccion.page_height = Inches(14)

    _agregar_membrete(doc)
    doc.add_paragraph()

    if tipo_doc == "Carta Documento (Intimación)":
        doc.add_heading('CARTA DOCUMENTO', 1)
        doc.add_paragraph(f"REMITENTE: {datos['Remitente (Nombre y DNI)']}")
        doc.add_paragraph(f"DESTINATARIO: {datos['Destinatario']}")
        doc.add_paragraph(f"DOMICILIO: {datos['Domicilio Destinatario']}, {datos['Localidad']}")
        doc.add_paragraph("-" * 40)
        p = doc.add_paragraph(datos['Texto de la intimación'])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph("\nQUEDA USTED DEBIDAMENTE INTIMADO Y NOTIFICADO.")
        doc.add_paragraph(f"\nFirma: ________________________\nAclaración: {datos['Remitente (Nombre y DNI)']}")

    elif tipo_doc == "Telegrama Laboral":
        doc.add_heading('TELEGRAMA LEY 23.789', 1)
        doc.add_paragraph(f"REMITENTE: {datos['Remitente']} (CUIL: {datos['CUIL Remitente']})")
        doc.add_paragraph(f"DESTINATARIO: {datos['Destinatario (Empleador)']}")
        doc.add_paragraph(f"DOMICILIO: {datos['Domicilio Laboral']}")
        doc.add_paragraph("-" * 40)
        p = doc.add_paragraph(datos['Texto del telegrama'])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph(f"\nFirma: ________________________\nAclaración: {datos['Remitente']}")

    elif tipo_doc == "Escrito: Conformidad (Sucesión)":
        doc.add_paragraph(f"Expte. N° {datos['Número de Expedte.']} - \"{datos['Carátula del Expediente']}\"").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph("\nPRESTA CONFORMIDAD. SE TENGA PRESENTE.\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run("SEÑOR JUEZ:\n").bold = True

        texto_cuerpo = (
            f"          {datos['Nombre del Cliente']}, DNI N° {datos['DNI/CUIT del Cliente']}, "
            f"por derecho propio, con domicilio legal constituido, en los autos caratulados "
            f"\"{datos['Carátula del Expediente']}\" (Expte. N° {datos['Número de Expedte.']}), "
            f"de trámite por ante este {datos['Radicación (Ej: Juzgado Civil y Com. 1 - Rosario)']}, "
            f"a V.S. respetuosamente digo:\n\n"
            f"          I. Que vengo por el presente acto a prestar mi expresa, total e irrevocable conformidad "
            f"con las operaciones de inventario, avalúo y partición presentadas en autos, solicitando se aprueben "
            f"en su totalidad y se dicte la correspondiente declaratoria/adjudicación.\n\n"
            f"          II. Que solicito se tenga presente lo expuesto para su oportunidad.\n\n"
            f"Proveer de conformidad,\nSERÁ JUSTICIA."
        )
        p_cuerpo = doc.add_paragraph(texto_cuerpo)
        p_cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    elif tipo_doc == "Contrato de Locación":
        doc.add_heading('CONTRATO DE LOCACIÓN DE INMUEBLE', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

        texto_contrato = (
            f"Entre el/la Sr/a. {datos['Propietario']}, DNI {datos['DNI Propietario']}, en adelante denominado "
            f"el 'LOCADOR', y el/la Sr/a. {datos['Inquilino']}, DNI {datos['DNI Inquilino']}, en adelante denominado "
            f"el 'LOCATARIO', se conviene celebrar el presente contrato de locación sujeto a las siguientes cláusulas:\n\n"
            f"PRIMERA (Objeto): El LOCADOR da en locación al LOCATARIO un inmueble sito en {datos['Dirección del Inmueble']}. "
            f"El inmueble se entrega con el siguiente inventario y estado: {datos['Inventario y estado']}.\n\n"
            f"SEGUNDA (Plazo): El plazo de la locación se fija en {datos['Plazo en meses']} meses a partir de la firma del presente.\n\n"
            f"TERCERA (Precio): El canon locativo se pacta en la suma de $ {datos['Precio Mensual ($)']} mensuales.\n\n"
            f"CUARTA (Servicios): El pago de los servicios de {datos['Servicios a cargo de']} estarán a cargo exclusivo del LOCATARIO.\n\n"
            f"QUINTA (Garantías): Afianzando las obligaciones de este contrato, se presentan las siguientes garantías: {datos['Garantías (Detalle)']}.\n\n"
            f"En prueba de conformidad, se firman dos ejemplares de un mismo tenor y a un solo efecto, "
            f"a los {datetime.now().strftime('%d días del mes de %m del año %Y')}."
        )
        p = doc.add_paragraph(texto_contrato)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(ruta_guardado)


class _PestanaGenerador(tk.Frame):

    def __init__(self, parent):
        super().__init__(parent)

        self._clientes = []
        self._roles_actuales = []

        ttk.Label(self, text="⚖️ Redacción Automatizada de Documentos", font=("Arial", 14, "bold")).pack(pady=(15, 10))

        marco_selector = tk.Frame(self)
        marco_selector.pack(fill=tk.X, padx=30, pady=10)

        ttk.Label(marco_selector, text="1. Seleccione el tipo de documento:").pack(side=tk.LEFT, padx=(0, 10))
        self.combo_tipo = ttk.Combobox(marco_selector, values=list(PLANTILLAS.keys()), width=40, state="readonly")
        self.combo_tipo.pack(side=tk.LEFT)
        self.combo_tipo.bind("<<ComboboxSelected>>", self.actualizar_formulario)
        agregar_tooltip(self.combo_tipo, "Elegí qué tipo de documento generar; abajo aparecen los "
                                           "campos que hacen falta completar para esa plantilla.")

        # ---------------- autocompletar con un cliente existente ----------------

        self.marco_autocompletar = ttk.LabelFrame(self, text="Autocompletar con un cliente (opcional)", padding=10)
        self.marco_autocompletar.pack(fill=tk.X, padx=30, pady=(0, 10))

        fila_ac = ttk.Frame(self.marco_autocompletar)
        fila_ac.pack(fill=tk.X)

        ttk.Label(fila_ac, text="Buscar cliente:").pack(side=tk.LEFT, padx=(0, 8))
        self.entry_buscar_cliente = tk.Entry(fila_ac, width=28)
        self.entry_buscar_cliente.pack(side=tk.LEFT, padx=(0, 8))
        self.entry_buscar_cliente.bind("<KeyRelease>", lambda e: self._renderizar_clientes())
        agregar_tooltip(self.entry_buscar_cliente, "Filtra la lista de abajo en vivo, por nombre de cliente.")

        self.btn_refrescar_clientes = ttk.Button(fila_ac, text="Actualizar lista", command=self._cargar_clientes)
        self.btn_refrescar_clientes.pack(side=tk.LEFT)

        ttk.Label(fila_ac, text="Rol en el documento:").pack(side=tk.LEFT, padx=(15, 8))
        self.combo_rol = ttk.Combobox(fila_ac, width=28, state="readonly")
        self.combo_rol.pack(side=tk.LEFT, padx=(0, 8))
        agregar_tooltip(self.combo_rol, "En qué campo del documento va este cliente (ej. remitente, "
                                          "destinatario, propietario, inquilino).")

        self.btn_completar_cliente = ttk.Button(fila_ac, text="Completar campos", command=self._completar_con_cliente)
        self.btn_completar_cliente.pack(side=tk.LEFT)
        agregar_tooltip(self.btn_completar_cliente, "Copia el nombre (y DNI/domicilio si la ficha los "
                                                       "tiene) del cliente seleccionado al campo del rol "
                                                       "elegido arriba. El resto de los campos hay que "
                                                       "completarlos igual a mano.")

        self.tabla_clientes = ttk.Treeview(
            self.marco_autocompletar, columns=("cliente", "ciudad"), show="headings", height=4,
        )
        self.tabla_clientes.heading("cliente", text="Cliente")
        self.tabla_clientes.heading("ciudad", text="Ciudad")
        self.tabla_clientes.column("cliente", width=340, anchor="w")
        self.tabla_clientes.column("ciudad", width=150, anchor="w")
        self.tabla_clientes.pack(fill=tk.X, pady=(8, 0))
        agregar_tooltip(self.tabla_clientes, "Elegí el cliente que va a ir en el rol seleccionado arriba.")

        self.marco_formulario = tk.Frame(self, padx=20, pady=20)
        self.marco_formulario.pack(fill=tk.BOTH, expand=True, padx=30, pady=10)

        self.campos_entry = {}

        self.btn_generar = ttk.Button(self, text="📄 Generar Documento Word", command=self.procesar_y_guardar)
        agregar_tooltip(self.btn_generar, "Genera el documento Word con los datos que completaste y "
                                            "te pregunta dónde guardarlo.")

        self._cargar_clientes()

    # ---------------- autocompletar con un cliente existente ----------------

    def _cargar_clientes(self):
        self.btn_refrescar_clientes.config(state=tk.DISABLED)

        def trabajo():
            _, _, carpetas_clientes = obtener_clientes()

            def terminar():
                self._clientes = sorted(carpetas_clientes, key=lambda p: p.name.lower())
                self._renderizar_clientes()
                self.btn_refrescar_clientes.config(state=tk.NORMAL)

            self.after(0, terminar)

        threading.Thread(target=trabajo, daemon=True).start()

    def _renderizar_clientes(self):
        self.tabla_clientes.delete(*self.tabla_clientes.get_children())
        filtro = self.entry_buscar_cliente.get().strip().lower()
        for carpeta in self._clientes:
            if filtro and filtro not in carpeta.name.lower():
                continue
            self.tabla_clientes.insert("", tk.END, values=(carpeta.name, carpeta.parent.name), iid=str(carpeta))

    def _completar_con_cliente(self):
        seleccion = self.tabla_clientes.selection()
        if not seleccion:
            messagebox.showinfo("Nada seleccionado", "Elegí un cliente de la lista primero.")
            return

        rol_elegido = self.combo_rol.get()
        rol = next((r for r in self._roles_actuales if r["campo_nombre"] == rol_elegido), None)
        if not rol:
            messagebox.showinfo("Elegí un rol", "Elegí en qué campo del documento va este cliente.")
            return

        carpeta_cliente = Path(seleccion[0])
        ficha = leer_ficha(carpeta_cliente)
        datos = ficha["datos"] if ficha else {}

        nombre = (datos.get("Apellido y Nombre") or carpeta_cliente.name).strip()
        dni = (datos.get("DNI/CUIT") or "").strip()
        domicilio = (datos.get("Dirección") or "").strip()

        if rol.get("combinar_dni_en_nombre") and dni:
            nombre = f"{nombre} - DNI {dni}"

        def poner_valor(campo, valor):
            widget = self.campos_entry.get(campo)
            if widget is None or not valor:
                return
            if isinstance(widget, tk.Text):
                widget.delete("1.0", tk.END)
                widget.insert("1.0", valor)
            else:
                widget.delete(0, tk.END)
                widget.insert(0, valor)

        poner_valor(rol["campo_nombre"], nombre)
        poner_valor(rol.get("campo_dni"), dni)
        poner_valor(rol.get("campo_domicilio"), domicilio)

        if not ficha:
            messagebox.showinfo(
                "Cliente sin ficha",
                f"'{carpeta_cliente.name}' todavía no tiene ficha guardada -- se completó solo el "
                f"nombre (de la carpeta). El resto hay que completarlo a mano."
            )

    def actualizar_formulario(self, event=None):
        for widget in self.marco_formulario.winfo_children():
            widget.destroy()
        self.campos_entry.clear()

        tipo_seleccionado = self.combo_tipo.get()
        campos_necesarios = PLANTILLAS[tipo_seleccionado]

        self._roles_actuales = ROLES_PLANTILLA.get(tipo_seleccionado, [])
        self.combo_rol.config(values=[r["campo_nombre"] for r in self._roles_actuales])
        self.combo_rol.set("")

        ttk.Label(self.marco_formulario, text="2. Complete los datos para la plantilla:", font=("Arial", 12, "bold")).grid(
            row=0, column=0, columnspan=2, pady=(0, 15), sticky="w"
        )

        fila = 1
        for campo in campos_necesarios:
            ttk.Label(self.marco_formulario, text=f"{campo}:").grid(row=fila, column=0, sticky="e", padx=10, pady=5)

            if "Texto" in campo or "Inventario" in campo or "Garantías" in campo:
                caja_texto = tk.Text(self.marco_formulario, height=5, width=50, font=("Arial", 10))
                caja_texto.grid(row=fila, column=1, sticky="w", pady=5)
                self.campos_entry[campo] = caja_texto
            else:
                entrada = ttk.Entry(self.marco_formulario, width=50, font=("Arial", 10))
                entrada.grid(row=fila, column=1, sticky="w", pady=5)
                self.campos_entry[campo] = entrada
            fila += 1

        self.btn_generar.pack(pady=20, fill=tk.X, padx=150)

    def procesar_y_guardar(self):
        tipo_doc = self.combo_tipo.get()

        datos_recopilados = {}
        for campo, widget in self.campos_entry.items():
            if isinstance(widget, tk.Text):
                datos_recopilados[campo] = widget.get("1.0", tk.END).strip()
            else:
                datos_recopilados[campo] = widget.get().strip()

            if not datos_recopilados[campo]:
                messagebox.showwarning("Campos vacíos", f"Por favor, complete el campo: {campo}")
                return

        ruta_guardado = filedialog.asksaveasfilename(
            title="Guardar documento final como...",
            initialfile=f"Borrador_{tipo_doc.replace(' ', '_')}.docx",
            defaultextension=".docx",
            filetypes=[("Documento de Word", "*.docx")]
        )

        if ruta_guardado:
            try:
                generar_word(tipo_doc, datos_recopilados, ruta_guardado)
                messagebox.showinfo("¡Éxito!", f"El documento fue redactado automáticamente y guardado en:\n{ruta_guardado}")
            except Exception as e:
                messagebox.showerror("Error", f"Ocurrió un error al crear el Word:\n{e}")


def construir_pestana(parent):
    return _PestanaGenerador(parent)
