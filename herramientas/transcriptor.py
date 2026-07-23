import os
import threading
import tkinter as tk
from tkinter import filedialog, ttk
import tkinter.scrolledtext as scrolledtext
from docx import Document
import PyPDF2
from PIL import Image

# Intentamos cargar la librería de imágenes
try:
    import pytesseract
    TESSERACT_DISPONIBLE = True
except ImportError:
    TESSERACT_DISPONIBLE = False

def extraer_texto_pdf(ruta_archivo):
    texto_completo = ""
    try:
        with open(ruta_archivo, 'rb') as archivo:
            lector = PyPDF2.PdfReader(archivo)
            for pagina in lector.pages:
                texto_pagina = pagina.extract_text()
                if texto_pagina:
                    texto_completo += texto_pagina + "\n"
        return texto_completo.strip()
    except Exception as e:
        return f"Error al leer PDF: {e}"

def extraer_texto_imagen(ruta_archivo):
    if not TESSERACT_DISPONIBLE:
        return "ERROR_LIBRERIA"
    try:
        imagen = Image.open(ruta_archivo)
        # lang='spa' es para que reconozca ñ y acentos en español
        texto = pytesseract.image_to_string(imagen, lang='spa')
        return texto.strip()
    except Exception as e:
        return f"ERROR_TESSERACT: {e}"

def proceso_transcripcion(ruta_archivo, ruta_guardado, consola, ventana, btn_iniciar):
    nombre_archivo, extension = os.path.splitext(os.path.basename(ruta_archivo))
    extension = extension.lower()
    
    consola.insert(tk.END, f"► Analizando: '{nombre_archivo}{extension}'\n")
    consola.insert(tk.END, f"► Extrayendo texto (esto puede demorar unos segundos)...\n")
    ventana.update()

    texto_extraido = ""

    # Detección de formato
    if extension == '.pdf':
        texto_extraido = extraer_texto_pdf(ruta_archivo)
    elif extension in ['.jpg', '.jpeg', '.png']:
        texto_extraido = extraer_texto_imagen(ruta_archivo)
        
        # Manejo de errores específicos para imágenes
        if texto_extraido == "ERROR_LIBRERIA":
            consola.insert(tk.END, "\n❌ Falta la librería en Python. Ejecuta en terminal: pip3 install pytesseract Pillow\n")
            consola.insert(tk.END, "-" * 60 + "\n\n")
            btn_iniciar.config(state=tk.NORMAL)
            return
        elif str(texto_extraido).startswith("ERROR_TESSERACT"):
            consola.insert(tk.END, "\n❌ Tesseract no está instalado en tu Mac.\nDebes abrir la terminal y ejecutar: brew install tesseract tesseract-lang\n")
            consola.insert(tk.END, "-" * 60 + "\n\n")
            btn_iniciar.config(state=tk.NORMAL)
            return
    else:
        consola.insert(tk.END, f"\n❌ Formato {extension} no soportado.\n")
        btn_iniciar.config(state=tk.NORMAL)
        return

    # Guardado del Word
    if texto_extraido and not texto_extraido.startswith("Error"):
        try:
            doc = Document()
            doc.add_heading(f'Transcripción del documento', 0)
            doc.add_paragraph(texto_extraido)
            
            doc.save(ruta_guardado)
            
            consola.insert(tk.END, f"\n✅ TRANSCRIPCIÓN FINALIZADA CON ÉXITO\n")
            consola.insert(tk.END, f"📄 Guardado en:\n  {ruta_guardado}\n")
        except Exception as e:
            consola.insert(tk.END, f"\n❌ Error al guardar el archivo Word: {e}\n")
    else:
        consola.insert(tk.END, f"\n⚠️ ATENCIÓN: No se detectó texto.\nSi es un PDF escaneado, probá subirlo como imagen JPG.\n")

    consola.insert(tk.END, "-" * 60 + "\n\n")
    consola.see(tk.END)
    btn_iniciar.config(state=tk.NORMAL)

def iniciar_transcripcion(consola, ventana, btn_iniciar):
    # 1. Seleccionar archivo de origen (PDF o Imagen)
    ruta_archivo = filedialog.askopenfilename(
        title="Seleccionar documento a transcribir",
        filetypes=[("PDF e Imágenes", "*.pdf *.jpg *.jpeg *.png")]
    )
    
    if not ruta_archivo:
        return

    # 2. Elegir dónde guardar el Word (Guardar como...)
    nombre_base = os.path.splitext(os.path.basename(ruta_archivo))[0]
    ruta_guardado = filedialog.asksaveasfilename(
        title="Guardar Word transcrito como...",
        initialfile=f"{nombre_base}_TRANSCRITO.docx",
        defaultextension=".docx",
        filetypes=[("Documento de Word", "*.docx")]
    )

    if not ruta_guardado:
        return # Si el usuario cancela el guardado, se detiene el proceso

    # 3. Iniciar el proceso
    btn_iniciar.config(state=tk.DISABLED)
    consola.insert(tk.END, f"\n" + "="*60 + "\n")
    
    hilo = threading.Thread(target=proceso_transcripcion, args=(ruta_archivo, ruta_guardado, consola, ventana, btn_iniciar))
    hilo.start()

def main():
    ventana = tk.Tk()
    ventana.title("Sistema de Transcripción a Word")
    ventana.geometry("700x550")
    ventana.configure(bg="#2C3E50")

    style = ttk.Style()
    style.theme_use('clam')
    style.configure("TFrame", background="#2C3E50")
    style.configure("TButton", font=("Helvetica", 12, "bold"), padding=12, background="#34495E", foreground="white")
    style.map("TButton", background=[("active", "#415B76")])
    style.configure("Title.TLabel", font=("Helvetica", 16, "bold"), background="#2C3E50", foreground="#ECF0F1")
    style.configure("Sub.TLabel", font=("Helvetica", 11), background="#2C3E50", foreground="#BDC3C7")

    frame_superior = ttk.Frame(ventana)
    frame_superior.pack(pady=25, fill=tk.X)

    titulo = ttk.Label(frame_superior, text="⚖️ Transcriptor de PDF e Imágenes a Word", style="Title.TLabel")
    titulo.pack()

    subtitulo = ttk.Label(frame_superior, text="Seleccioná un archivo y elegí dónde guardar la versión editable.", style="Sub.TLabel")
    subtitulo.pack(pady=(5, 0))

    btn_iniciar = ttk.Button(ventana, text="📄 Seleccionar Archivo y Transcribir")
    btn_iniciar.config(command=lambda: iniciar_transcripcion(consola, ventana, btn_iniciar))
    btn_iniciar.pack(pady=10)

    frame_consola = tk.Frame(ventana, bg="#2C3E50")
    frame_consola.pack(fill=tk.BOTH, expand=True, padx=25, pady=(10, 25))
    
    consola = scrolledtext.ScrolledText(frame_consola, wrap=tk.WORD, font=("Consolas", 10), bg="#ECF0F1", fg="#2C3E50", bd=0, padx=15, pady=15)
    consola.pack(fill=tk.BOTH, expand=True)
    consola.insert(tk.END, "Sistema iniciado y listo para operar.\nEsperando selección de documento...\n\n")

    ventana.mainloop()

if __name__ == "__main__":
    main()
