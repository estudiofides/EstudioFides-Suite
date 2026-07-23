"""
Pestaña "Transcriptor" para la app unificada (hub_app.py).
Misma lógica de herramientas/transcriptor.py, embebida como Frame.
(Se saca el ventana.update() manual que hacía el original: llamarlo
desde el hilo de fondo no es seguro en Tkinter; no hace falta, los
inserts en la consola se ven igual.)
"""
import os
import threading
import tkinter as tk
from tkinter import filedialog, ttk
import tkinter.scrolledtext as scrolledtext
from docx import Document
import PyPDF2
from PIL import Image

from src.tooltip import agregar_tooltip

try:
    import pytesseract
    TESSERACT_DISPONIBLE = True
except ImportError:
    TESSERACT_DISPONIBLE = False


def extraer_texto_pdf(ruta_archivo):
    texto_completo = ""
    try:
        with open(ruta_archivo, 'rb') as archivo:
            lector = PyPDF2.PdfReader(archivo)
            for pagina in lector.pages:
                texto_pagina = pagina.extract_text()
                if texto_pagina:
                    texto_completo += texto_pagina + "\n"
        return texto_completo.strip()
    except Exception as e:
        return f"Error al leer PDF: {e}"


def extraer_texto_imagen(ruta_archivo):
    if not TESSERACT_DISPONIBLE:
        return "ERROR_LIBRERIA"
    try:
        imagen = Image.open(ruta_archivo)
        texto = pytesseract.image_to_string(imagen, lang='spa')
        return texto.strip()
    except Exception as e:
        return f"ERROR_TESSERACT: {e}"


def construir_pestana(parent):

    frame = ttk.Frame(parent, padding=15)

    ttk.Label(frame, text="⚖️ Transcriptor de PDF e Imágenes a Word", font=("Arial", 14, "bold")).pack(pady=(5, 5))
    ttk.Label(
        frame,
        text="Seleccioná un archivo y elegí dónde guardar la versión editable.",
        foreground="#555555"
    ).pack(pady=(0, 10))

    btn_iniciar = ttk.Button(frame, text="📄 Seleccionar Archivo y Transcribir")
    btn_iniciar.pack(pady=10)
    agregar_tooltip(btn_iniciar, "Elegí un PDF o imagen (JPG/PNG) y lo convierte a un Word editable. "
                                   "Si el PDF no tiene texto (escaneado), usa OCR.")

    consola = scrolledtext.ScrolledText(frame, wrap=tk.WORD, font=("Consolas", 10), bg="#f4f4f4")
    consola.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
    consola.insert(tk.END, "Sistema iniciado y listo para operar.\nEsperando selección de documento...\n\n")

    def proceso_transcripcion(ruta_archivo, ruta_guardado):

        nombre_archivo, extension = os.path.splitext(os.path.basename(ruta_archivo))
        extension = extension.lower()

        consola.insert(tk.END, f"► Analizando: '{nombre_archivo}{extension}'\n")
        consola.insert(tk.END, f"► Extrayendo texto (esto puede demorar unos segundos)...\n")
        consola.see(tk.END)

        texto_extraido = ""

        if extension == '.pdf':
            texto_extraido = extraer_texto_pdf(ruta_archivo)
        elif extension in ['.jpg', '.jpeg', '.png']:
            texto_extraido = extraer_texto_imagen(ruta_archivo)

            if texto_extraido == "ERROR_LIBRERIA":
                consola.insert(tk.END, "\n❌ Falta la librería en Python. Ejecuta en terminal: pip3 install pytesseract Pillow --break-system-packages\n")
                consola.insert(tk.END, "-" * 60 + "\n\n")
                btn_iniciar.config(state=tk.NORMAL)
                return
            elif str(texto_extraido).startswith("ERROR_TESSERACT"):
                consola.insert(tk.END, "\n❌ Tesseract no está instalado en tu Mac.\nDebes abrir la terminal y ejecutar: brew install tesseract tesseract-lang\n")
                consola.insert(tk.END, "-" * 60 + "\n\n")
                btn_iniciar.config(state=tk.NORMAL)
                return
        else:
            consola.insert(tk.END, f"\n❌ Formato {extension} no soportado.\n")
            btn_iniciar.config(state=tk.NORMAL)
            return

        if texto_extraido and not texto_extraido.startswith("Error"):
            try:
                doc = Document()
                doc.add_heading('Transcripción del documento', 0)
                doc.add_paragraph(texto_extraido)

                doc.save(ruta_guardado)

                consola.insert(tk.END, f"\n✅ TRANSCRIPCIÓN FINALIZADA CON ÉXITO\n")
                consola.insert(tk.END, f"📄 Guardado en:\n  {ruta_guardado}\n")
            except Exception as e:
                consola.insert(tk.END, f"\n❌ Error al guardar el archivo Word: {e}\n")
        else:
            consola.insert(tk.END, f"\n⚠️ ATENCIÓN: No se detectó texto.\nSi es un PDF escaneado, probá subirlo como imagen JPG.\n")

        consola.insert(tk.END, "-" * 60 + "\n\n")
        consola.see(tk.END)
        btn_iniciar.config(state=tk.NORMAL)

    def iniciar_transcripcion():

        ruta_archivo = filedialog.askopenfilename(
            title="Seleccionar documento a transcribir",
            filetypes=[("PDF e Imágenes", "*.pdf *.jpg *.jpeg *.png")]
        )
        if not ruta_archivo:
            return

        nombre_base = os.path.splitext(os.path.basename(ruta_archivo))[0]
        ruta_guardado = filedialog.asksaveasfilename(
            title="Guardar Word transcrito como...",
            initialfile=f"{nombre_base}_TRANSCRITO.docx",
            defaultextension=".docx",
            filetypes=[("Documento de Word", "*.docx")]
        )
        if not ruta_guardado:
            return

        btn_iniciar.config(state=tk.DISABLED)
        consola.insert(tk.END, "\n" + "=" * 60 + "\n")

        hilo = threading.Thread(
            target=proceso_transcripcion,
            args=(ruta_archivo, ruta_guardado),
            daemon=True
        )
        hilo.start()

    btn_iniciar.config(command=iniciar_transcripcion)

    return frame
